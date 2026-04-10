import re
from datetime import datetime
from pathlib import Path
from typing import Optional

from src.config import EXPORTS_DIR, EXPORT, PDF
from src.core.agents.analyzer_agent import MeetingData


class GeneratorAgent:
    def __init__(self, dossier_export: Optional[Path] = None):
        self._dossier = Path(dossier_export) if dossier_export else EXPORTS_DIR
        self._dossier.mkdir(parents=True, exist_ok=True)

    def generer(
        self,
        markdown: str,
        donnees: MeetingData,
        formats: Optional[list] = None,
    ) -> dict:
        formats = formats or ["pdf", "word", "txt"]
        nom_base = self._nom_fichier(donnees)
        fichiers = {}

        generateurs = {
            "txt": self._exporter_txt,
            "word": self._exporter_word,
            "pdf": self._exporter_pdf,
        }

        for fmt in formats:
            if fmt not in generateurs:
                continue
            ext = "docx" if fmt == "word" else fmt
            chemin = self._dossier / f"{nom_base}.{ext}"
            try:
                generateurs[fmt](chemin, markdown, donnees)
                fichiers[fmt] = str(chemin)
            except Exception as e:
                fichiers[f"{fmt}_error"] = str(e)

        return fichiers

    def _nom_fichier(self, donnees: MeetingData) -> str:
        titre = re.sub(r"[^\w\s-]", "", donnees.titre or "reunion")
        titre = re.sub(r"\s+", "_", titre.strip())[: EXPORT["titre_max_chars"]]
        horodatage = datetime.now().strftime("%Y%m%d_%H%M")
        return f"CR_{titre}_{horodatage}"

    # ── TXT

    def _exporter_txt(self, chemin: Path, markdown: str, donnees: MeetingData):
        L = EXPORT["txt_largeur"]
        lignes = []

        def sep(c="═"):
            lignes.append(c * L)

        def titre_sec(t):
            lignes.extend(["", "─" * L, f"  {t}", "─" * L])

        sep()
        lignes.append("  COMPTE RENDU DE RÉUNION".center(L))
        lignes.append(f"  {donnees.titre or ''}".center(L))
        sep()

        for label, val in [
            ("Date", donnees.date),
            ("Lieu", donnees.lieu),
            ("Organisateur", donnees.organisateur),
            ("Durée", donnees.duree),
            ("Participants", ", ".join(donnees.participants)),
        ]:
            if val:
                lignes.append(f"  {label:<15}: {val}")

        if donnees.ordre_du_jour:
            titre_sec("ORDRE DU JOUR")
            for i, p in enumerate(donnees.ordre_du_jour, 1):
                lignes.append(f"  {i}. {p}")

        if donnees.resume:
            titre_sec("RÉSUMÉ")
            lignes.append(f"  {donnees.resume}")

        if donnees.points_discutes:
            titre_sec("POINTS DISCUTÉS")
            for p in donnees.points_discutes:
                lignes.append(f"  • {p}")

        if donnees.decisions:
            titre_sec("DÉCISIONS PRISES")
            for d in donnees.decisions:
                lignes.append(f"  [OK] {d}")

        if donnees.actions:
            titre_sec("ACTIONS À MENER")
            for i, a in enumerate(donnees.actions, 1):
                ech = f" | {a.echeance}" if a.echeance else ""
                lignes.append(f"  {i}. {a.texte}")
                lignes.append(f"     → {a.responsable}{ech} [{a.priorite.upper()}]")

        if donnees.prochaines_etapes:
            titre_sec("PROCHAINES ÉTAPES")
            for i, e in enumerate(donnees.prochaines_etapes, 1):
                lignes.append(f"  {i}. {e}")

        sep()
        lignes.append(
            f"  Généré par MeetScribe le {datetime.now().strftime('%d/%m/%Y à %H:%M')}".center(
                L
            )
        )
        sep()

        chemin.write_text("\n".join(lignes), encoding="utf-8")

    # ── WORD

    def _exporter_word(self, chemin: Path, markdown: str, donnees: MeetingData):
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(donnees.titre or "Compte rendu de réunion")
        run.bold = True
        run.font.size = Pt(18)

        for label, val in [
            ("Date", donnees.date),
            ("Lieu", donnees.lieu),
            ("Organisateur", donnees.organisateur),
            ("Participants", ", ".join(donnees.participants)),
        ]:
            if val:
                par = doc.add_paragraph()
                par.add_run(f"{label} : ").bold = True
                par.add_run(val)

        doc.add_paragraph()

        def section(titre_sec, items):
            if not items:
                return
            doc.add_heading(titre_sec, level=2)
            for item in items:
                doc.add_paragraph(item)

        section(
            "Ordre du jour",
            [f"{i}. {p}" for i, p in enumerate(donnees.ordre_du_jour, 1)],
        )

        if donnees.resume:
            doc.add_heading("Résumé", level=2)
            doc.add_paragraph(donnees.resume)

        section("Points discutés", [f"• {p}" for p in donnees.points_discutes])
        section("Décisions prises", [f"• {d}" for d in donnees.decisions])

        if donnees.actions:
            doc.add_heading("Actions à mener", level=2)
            table = doc.add_table(rows=1, cols=4)
            table.style = "Table Grid"
            for i, h in enumerate(["Action", "Responsable", "Échéance", "Priorité"]):
                cell = table.rows[0].cells[i]
                cell.text = h
                cell.paragraphs[0].runs[0].bold = True
            for a in donnees.actions:
                row = table.add_row()
                for i, val in enumerate(
                    [a.texte, a.responsable, a.echeance or "—", a.priorite.upper()]
                ):
                    row.cells[i].text = val

        section(
            "Prochaines étapes",
            [f"{i}. {e}" for i, e in enumerate(donnees.prochaines_etapes, 1)],
        )

        doc.add_paragraph()
        pied = doc.add_paragraph()
        pied.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pied.add_run(
            f"Généré par MeetScribe le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
        ).italic = True

        doc.save(str(chemin))

    # ── PDF

    def _exporter_pdf(self, chemin: Path, markdown: str, donnees: MeetingData):
        from reportlab.lib.pagesizes import A4
        from reportlab.lib import colors
        from reportlab.lib.units import cm
        from reportlab.platypus import (
            SimpleDocTemplate,
            Paragraph,
            Spacer,
            Table,
            TableStyle,
            HRFlowable,
        )
        from reportlab.lib.styles import ParagraphStyle
        from reportlab.lib.enums import TA_CENTER

        def c(key):
            return colors.HexColor(f"#{PDF[key]}")

        s_titre = ParagraphStyle(
            "titre",
            fontName="Helvetica-Bold",
            fontSize=20,
            textColor=c("couleur_accent"),
            alignment=TA_CENTER,
            spaceAfter=6,
        )
        s_meta = ParagraphStyle(
            "meta",
            fontName="Helvetica",
            fontSize=9,
            textColor=c("couleur_discret"),
            alignment=TA_CENTER,
            spaceAfter=4,
        )
        s_section = ParagraphStyle(
            "section",
            fontName="Helvetica-Bold",
            fontSize=12,
            textColor=c("couleur_titre"),
            spaceBefore=14,
            spaceAfter=6,
        )
        s_corps = ParagraphStyle(
            "corps",
            fontName="Helvetica",
            fontSize=10,
            textColor=c("couleur_texte"),
            leading=15,
            spaceAfter=4,
        )
        s_bullet = ParagraphStyle(
            "bullet",
            fontName="Helvetica",
            fontSize=10,
            textColor=c("couleur_texte"),
            leftIndent=14,
            leading=15,
            spaceAfter=3,
        )

        def bg(canvas, doc):
            canvas.saveState()
            canvas.setFillColor(c("couleur_fond"))
            canvas.rect(0, 0, A4[0], A4[1], fill=1, stroke=0)
            canvas.setFillColor(c("couleur_accent"))
            canvas.rect(0, A4[1] - 5, A4[0], 5, fill=1, stroke=0)
            canvas.setFont("Helvetica", 8)
            canvas.setFillColor(c("couleur_discret"))
            canvas.drawCentredString(
                A4[0] / 2,
                20,
                f"Généré par MeetScribe le {datetime.now().strftime('%d/%m/%Y %H:%M')}"
                f"  |  Page {doc.page}",
            )
            canvas.restoreState()

        doc_pdf = SimpleDocTemplate(
            str(chemin),
            pagesize=A4,
            leftMargin=1.8 * cm,
            rightMargin=1.8 * cm,
            topMargin=1.8 * cm,
            bottomMargin=1.5 * cm,
        )
        story = []

        story.append(Paragraph("COMPTE RENDU DE RÉUNION", s_meta))
        story.append(Paragraph(donnees.titre or "Réunion", s_titre))

        meta = [
            v
            for v in [
                donnees.date,
                donnees.lieu,
                donnees.organisateur,
                ", ".join(donnees.participants[:5]),
            ]
            if v
        ]
        if meta:
            story.append(Paragraph("  |  ".join(meta), s_meta))

        story.append(
            HRFlowable(
                width="100%",
                thickness=1,
                color=c("couleur_grille"),
                spaceAfter=10,
            )
        )

        def ajouter_section(titre_sec, items, prefix="→"):
            if not items:
                return
            story.append(Paragraph(titre_sec, s_section))
            for item in items:
                story.append(Paragraph(f"{prefix}  {item}", s_bullet))

        ajouter_section("Ordre du jour", donnees.ordre_du_jour)

        if donnees.resume:
            story.append(Paragraph("Résumé", s_section))
            story.append(Paragraph(donnees.resume, s_corps))

        ajouter_section("Points discutés", donnees.points_discutes)
        ajouter_section("Décisions prises", donnees.decisions, prefix="✓")

        if donnees.actions:
            story.append(Paragraph("Actions à mener", s_section))
            story.append(Spacer(1, 6))
            col_w = [
                doc_pdf.width * 0.44,
                doc_pdf.width * 0.20,
                doc_pdf.width * 0.20,
                doc_pdf.width * 0.16,
            ]
            data = [["Action", "Responsable", "Échéance", "Priorité"]]
            for a in donnees.actions:
                data.append(
                    [a.texte, a.responsable, a.echeance or "—", a.priorite.upper()]
                )

            prio_c = {
                "HAUTE": PDF["couleur_rouge"],
                "NORMALE": PDF["couleur_accent"],
                "BASSE": PDF["couleur_violet"],
            }
            t_style = [
                ("BACKGROUND", (0, 0), (-1, 0), c("couleur_surface")),
                ("TEXTCOLOR", (0, 0), (-1, 0), c("couleur_titre")),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 9),
                ("TEXTCOLOR", (0, 1), (-1, -1), c("couleur_texte")),
                ("GRID", (0, 0), (-1, -1), 0.5, c("couleur_grille")),
                (
                    "ROWBACKGROUNDS",
                    (0, 1),
                    (-1, -1),
                    [c("couleur_fond"), c("couleur_surface")],
                ),
                ("TOPPADDING", (0, 0), (-1, -1), 6),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ]
            for i, a in enumerate(donnees.actions, 1):
                couleur = colors.HexColor(
                    f"#{prio_c.get(a.priorite.upper(), PDF['couleur_accent'])}"
                )
                t_style.append(("TEXTCOLOR", (3, i), (3, i), couleur))
                t_style.append(("FONTNAME", (3, i), (3, i), "Helvetica-Bold"))

            table = Table(data, colWidths=col_w, repeatRows=1)
            table.setStyle(TableStyle(t_style))
            story.append(table)

        ajouter_section("Prochaines étapes", donnees.prochaines_etapes)

        doc_pdf.build(story, onFirstPage=bg, onLaterPages=bg)
